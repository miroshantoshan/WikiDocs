import os
from pathlib import Path
from time import sleep
import re
import textwrap

import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import set_styles, download_picture, clear_dir,clearing_tags, create_document_folder, create_pictures_folder
from config import user_agent,picture_folder_path,document_path,document_folder_path


load_dotenv()


RED = '\033[91m'
GREEN = '\033[92m'
YELLOW = '\033[93m'
BLUE = '\033[94m'
MAGENTA = '\033[95m'
GRAY = '\033[90m'
CYAN = '\033[96m'
WHITE = '\033[97m'
RESET = '\033[0m' 
BOLD = '\033[1m'

print(f'''

{CYAN}
╔══════════════════════════════════════════════════════════╗
║                                                          ║
║   ██╗    ██╗██╗██╗  ██╗██╗                               ║
║   ██║    ██║██║██║ ██╔╝██║                               ║
║   ██║ █╗ ██║██║█████╔╝ ██║                               ║
║   ██║███╗██║██║██╔═██╗ ██║                               ║
║   ╚███╔███╔╝██║██║  ██╗██║                               ║
║    ╚══╝╚══╝ ╚═╝╚═╝  ╚═╝╚═╝                               ║
║   {BLUE}                                                       ║
║   {BLUE}██████╗  ██████╗  ██████╗ ███████╗                     ║
║   {BLUE}██╔══██╗██╔═══██╗██╔════╝ ██╔════╝                     ║
║   {BLUE}██║  ██║██║   ██║██║      ███████╗                     ║
║   {BLUE}██║  ██║██║   ██║██║      ╚════██║                     ║
║   {BLUE}██████╔╝╚██████╔╝╚██████╔╝███████║                     ║
║   {BLUE}╚═════╝  ╚═════╝  ╚═════╝ ╚══════╝                     ║
║                                                          ║
╚══════════════════════════════════════════════════════════╝

{BOLD}{CYAN}This is a script for creating a document from a Wikipedia page. {BLUE}Please, choose your {RED}language:
    {RED}1. {GREEN}Русский
    {RED}2. {GREEN}English''')

start_menu = input(f'''{BLUE}-------------------------------------
{CYAN}Your choice: {RED}''')

if start_menu == '1':
    print(f'''







{GREEN}Вы выбрали руский язык. {CYAN}Пожалуйста, выберите вариант создания {BLUE}Word-документа:

          {RED}1. {CYAN}С помощью ссылки на статью в Википедии
          {RED}2. {CYAN}С помощью ключевого слова для статьи в Википедии''')
    choose_method_ru = input(f'''
{CYAN}----------------------------------------------

{BLUE}Your choose: {RED}''')
    if choose_method_ru == '1':
        url = input(f'''{BLUE}
--------------------------------------
{GREEN}Введите ссылку статьи с Википедии: {CYAN}''')
        print(f'''{CYAN}--------------------------------
{BLUE}Cоздание документа со статьи: {url}
{CYAN}----------------------------------------------------{GREEN}''')


    elif choose_method_ru == '2':
        keyword = input(f'''
{CYAN}-------------------------------------------------------
{GREEN}Введите ключевое слово: {CYAN}''')
        print(f'''{CYAN}---------------------------------------
{BLUE}Cоздание документа с ключевым словом: {keyword}
{CYAN}---------------------------------------------------------{GREEN}''')
        url = f'https://ru.wikipedia.org/wiki/{keyword}'



elif start_menu == '2':
    print(f'''







{GREEN}You have choosen English language. {CYAN}Please,choose the mode for {BLUE}creating a Word document:

          {RED}1. {CYAN}With a link to Wikipedia article
          {RED}2. {CYAN}With a keyword for Wikipedia article''')
    choose_method_ru = input(f'''
{CYAN}----------------------------------------------

{BLUE}Your choose: {RED}''')
    if choose_method_ru == '1':
        url = input(f'''{BLUE}
--------------------------------------
{GREEN}Enter the link to the article: {CYAN}''')
        print(f'''{CYAN}----------------------------------
{BLUE}Creating a document from article: {url}
{CYAN}----------------------------------------------------{GREEN}''')


    elif choose_method_ru == '2':
        keyword = input(f'''
{CYAN}-------------------------------------------------------
{GREEN}Еnter the article's keyword: {CYAN}''')
        print(f'''{CYAN}---------------------------------------
{BLUE}Creating a document with keyword: {keyword}
{CYAN}---------------------------------------------------------{GREEN}''')
        url = f'https://en.wikipedia.org/wiki/{keyword}'




create_pictures_folder(picture_folder_path)
create_document_folder(document_folder_path)

headers = user_agent

response = requests.get(url,verify=False, headers=headers)
response.raise_for_status()  

soup = BeautifulSoup(response.text, features='html.parser')

header = soup.find('h1', class_='firstHeading mw-first-heading')

content = soup.find('div', class_='mw-content-ltr mw-parser-output')

tags = content.find_all(['p','h2','img'])


doc = Document()

set_styles(doc,'Times New Roman',14)

head = doc.add_heading(header.text)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER


needed_tags = clearing_tags(tags)

for index,tag in enumerate(needed_tags):

    if "</p>" in str(tag):

        match = re.sub(r"\[\w*\]","", tag.text)
        paragraphs = doc.add_paragraph(match)
        paragraphs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if "<img" in str(tag):
        link_photo = f"https:{tag['src']}"
        picture_name = f"image{index}"
        path = Path(picture_folder_path, f"{picture_name}.png")

        download_picture(link_photo,path,headers)
        
        doc.add_picture(f"{picture_folder_path}/{picture_name}.png")
        picture_paragraph = doc.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "<h2" in str(tag):
        head2 = doc.add_heading(tag.text, level=2)
        head2.alignment = WD_ALIGN_PARAGRAPH.CENTER


doc.save(document_path)

if start_menu == '1':
    print(f'''
{BLUE}--------------------------------------------
{GREEN}Документ успешно сохранён в {CYAN}{document_path}
{BLUE}---------------------------------------------''')
else:
    print(f'''
{BLUE}--------------------------------------------
{GREEN}Document successfully saved in {CYAN}{document_path}
{BLUE}---------------------------------------------''')


clear_dir(picture_folder_path)

"*th*r*"


